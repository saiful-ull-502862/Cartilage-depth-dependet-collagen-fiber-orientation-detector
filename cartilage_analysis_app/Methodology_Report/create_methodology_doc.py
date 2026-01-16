"""
Script to generate Methodology document with equations in Word and PDF format.
Uses python-docx for Word document creation with proper equation support.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree

# Output directory
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def add_math_element(paragraph, latex_text):
    """
    Add an OMML (Office Math Markup Language) equation to a paragraph.
    This creates MathType-compatible equations in Word.
    """
    # Create the math namespace
    m_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    
    # Create oMath element
    oMath = OxmlElement('m:oMath')
    
    # Create run element
    r = OxmlElement('m:r')
    
    # Create text element
    t = OxmlElement('m:t')
    t.text = latex_text
    
    r.append(t)
    oMath.append(r)
    
    paragraph._p.append(oMath)

def create_equation_run(paragraph, equation_parts):
    """
    Create inline equation with proper formatting.
    equation_parts is a list of tuples: ('text', is_italic, is_subscript)
    """
    for part, is_italic, is_subscript in equation_parts:
        run = paragraph.add_run(part)
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
        run.italic = is_italic
        if is_subscript:
            run.font.subscript = True

def create_methodology_document():
    """Create the methodology Word document with equations."""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Add Title
    title = doc.add_heading('Methodology: Collagen Fibril Orientation Analysis', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Image Processing and Quantitative Analysis of PLM Images')
    run.italic = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacer
    
    # === PARAGRAPH 1: Image Processing Pipeline ===
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # First sentence
    p1.add_run('Collagen fibril orientation was quantified from Polarized Light Microscopy (PLM) images through a systematic image processing pipeline based on HSV (Hue, Saturation, Value) color space analysis. ')
    
    # Technical description
    p1.add_run('Each PLM image was first converted from BGR to HSV color space using OpenCV, where the Hue channel (0–179° in OpenCV convention) encodes the birefringence-dependent color information directly correlated with collagen fibril orientation. To eliminate background noise and non-birefringent regions, a dual-threshold mask was applied, excluding pixels with saturation (')
    
    # S < 10 equation
    run_s = p1.add_run('S')
    run_s.italic = True
    run_s.font.name = 'Cambria Math'
    p1.add_run(' < 10) or value (')
    run_v = p1.add_run('V')
    run_v.italic = True
    run_v.font.name = 'Cambria Math'
    p1.add_run(' < 10). ')
    
    # Angle calculation description
    p1.add_run('The fibril orientation angle (')
    run_theta = p1.add_run('θ')
    run_theta.italic = True
    run_theta.font.name = 'Cambria Math'
    p1.add_run(', ranging from 0° to 90°) was calculated through linear interpolation of the hue values, where the hue corresponding to the Superficial Zone (SZ) was mapped to 0° (fibrils parallel to the articular surface) and the hue corresponding to the Deep Zone (DZ) was mapped to 90° (fibrils perpendicular to the tidemark), following the equation:')
    
    # === Main Equation (centered) ===
    eq_para = doc.add_paragraph()
    eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_para.paragraph_format.space_before = Pt(12)
    eq_para.paragraph_format.space_after = Pt(12)
    
    # Build equation: θ = [(H − H₀) / (H₉₀ − H₀)] × 90°
    run = eq_para.add_run('θ')
    run.italic = True
    run.font.name = 'Cambria Math'
    run.font.size = Pt(14)
    
    eq_para.add_run(' = [(').font.size = Pt(14)
    
    run_h = eq_para.add_run('H')
    run_h.italic = True
    run_h.font.name = 'Cambria Math'
    run_h.font.size = Pt(14)
    
    eq_para.add_run(' − ').font.size = Pt(14)
    
    run_h0 = eq_para.add_run('H')
    run_h0.italic = True
    run_h0.font.name = 'Cambria Math'
    run_h0.font.size = Pt(14)
    run_sub0 = eq_para.add_run('₀')
    run_sub0.font.size = Pt(14)
    
    eq_para.add_run(') / (').font.size = Pt(14)
    
    run_h90 = eq_para.add_run('H')
    run_h90.italic = True
    run_h90.font.name = 'Cambria Math'
    run_h90.font.size = Pt(14)
    run_sub90 = eq_para.add_run('₉₀')
    run_sub90.font.size = Pt(14)
    
    eq_para.add_run(' − ').font.size = Pt(14)
    
    run_h02 = eq_para.add_run('H')
    run_h02.italic = True
    run_h02.font.name = 'Cambria Math'
    run_h02.font.size = Pt(14)
    run_sub02 = eq_para.add_run('₀')
    run_sub02.font.size = Pt(14)
    
    eq_para.add_run(')] × 90°').font.size = Pt(14)
    
    # Equation number
    eq_para.add_run('          (1)').font.size = Pt(12)
    
    # === Variable definitions ===
    var_para = doc.add_paragraph()
    var_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    var_para.add_run('where ')
    
    run_h = var_para.add_run('H')
    run_h.italic = True
    run_h.font.name = 'Cambria Math'
    var_para.add_run(' is the measured hue value, ')
    
    run_h0 = var_para.add_run('H')
    run_h0.italic = True
    run_h0.font.name = 'Cambria Math'
    var_para.add_run('₀ is the reference hue for SZ (0° orientation), and ')
    
    run_h90 = var_para.add_run('H')
    run_h90.italic = True
    run_h90.font.name = 'Cambria Math'
    var_para.add_run('₉₀ is the reference hue for DZ (90° orientation). ')
    
    # Continue paragraph
    var_para.add_run('To ensure robust hue averaging across the circular color space, circular mean calculations were employed using trigonometric functions to prevent artifacts at hue boundaries. The cartilage depth was normalized to tissue thickness and discretized into 100 equally spaced bins, with each bin analyzed independently to generate continuous depth-dependent orientation profiles. For each depth bin, the mean orientation angle (')
    
    run_theta_bar = var_para.add_run('θ̄')
    run_theta_bar.italic = True
    run_theta_bar.font.name = 'Cambria Math'
    
    var_para.add_run('), standard deviation (')
    
    run_sigma = var_para.add_run('σ')
    run_sigma.italic = True
    run_sigma.font.name = 'Cambria Math'
    
    var_para.add_run('), and representative color properties (RGB values, hex color, and intensity) were computed from all valid pixels within the corresponding horizontal slice.')
    
    doc.add_paragraph()  # Spacer
    
    # === PARAGRAPH 2: Zone Segmentation and Comparative Analysis ===
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p2.add_run('The articular cartilage was segmented into three anatomically distinct zones based on normalized tissue thickness: the ')
    p2.add_run('Superficial Zone').bold = True
    p2.add_run(' (SZ, 0–10%), the ')
    p2.add_run('Middle Zone').bold = True
    p2.add_run(' (MZ, 10–40%), and the ')
    p2.add_run('Deep Zone').bold = True
    p2.add_run(' (DZ, 40–100%). ')
    
    p2.add_run('To establish baseline orientation profiles for comparative analysis, one representative PLM image was selected as a reference for each healthy cartilage group—specifically, healthy Femoral Cartilage and healthy Tibial Cartilage. These reference images underwent the complete orientation analysis pipeline to establish the characteristic depth-dependent fibril angle profiles (')
    
    run_theta = p2.add_run('θ')
    run_theta.italic = True
    run_theta.font.name = 'Cambria Math'
    
    p2.add_run(' vs. normalized depth) for healthy tissue. Subsequently, the fibril orientation profiles of degraded cartilage specimens, categorized by disease progression as ')
    p2.add_run('Early Osteoarthritis (EOA)').bold = True
    p2.add_run(', ')
    p2.add_run('Moderate Osteoarthritis (MOA)').bold = True
    p2.add_run(', and ')
    p2.add_run('Advanced Osteoarthritis (AOA)').bold = True
    p2.add_run(', were analyzed using identical processing parameters (zone boundaries and color calibration) derived from their respective healthy references. ')
    
    p2.add_run('For both femoral and tibial compartments, the degraded group profiles were directly overlaid and statistically compared against the healthy reference profile via depth-wise mean angle calculations and standard deviation propagation. This reference-based comparative approach enabled quantitative assessment of zone-specific alterations in collagen fibril architecture—including disorganization, loss of surface tangential orientation, and disruption of the characteristic arcade-like structure—as a function of osteoarthritis severity progression.')
    
    # Save Word document
    docx_path = os.path.join(OUTPUT_DIR, 'Methodology_Fibril_Orientation_Analysis.docx')
    doc.save(docx_path)
    print(f"Word document saved: {docx_path}")
    
    return docx_path

def convert_to_pdf(docx_path):
    """Convert Word document to PDF using different methods."""
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    # Method 1: Try using docx2pdf (requires MS Word installed)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"PDF saved using docx2pdf: {pdf_path}")
        return pdf_path
    except ImportError:
        print("docx2pdf not installed. Trying alternative methods...")
    except Exception as e:
        print(f"docx2pdf failed: {e}. Trying alternative methods...")
    
    # Method 2: Try using comtypes (Windows with MS Word)
    try:
        import comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
        doc.Close()
        word.Quit()
        print(f"PDF saved using COM automation: {pdf_path}")
        return pdf_path
    except ImportError:
        print("comtypes not installed. Trying alternative methods...")
    except Exception as e:
        print(f"COM automation failed: {e}. Trying alternative methods...")
    
    # Method 3: Try using win32com (Windows with MS Word)
    try:
        import win32com.client
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        print(f"PDF saved using win32com: {pdf_path}")
        return pdf_path
    except ImportError:
        print("pywin32 not installed.")
    except Exception as e:
        print(f"win32com failed: {e}")
    
    print("\n" + "="*60)
    print("MANUAL CONVERSION REQUIRED")
    print("="*60)
    print(f"Word document created successfully: {docx_path}")
    print("\nTo convert to PDF:")
    print("1. Open the Word document")
    print("2. Go to File > Save As")
    print("3. Select 'PDF' as the file type")
    print("4. Click Save")
    print("="*60)
    
    return None

if __name__ == '__main__':
    print("Creating Methodology Document...")
    print("="*60)
    
    # Create Word document
    docx_path = create_methodology_document()
    
    # Convert to PDF
    pdf_path = convert_to_pdf(docx_path)
    
    print("\n" + "="*60)
    print("COMPLETE!")
    print(f"Word Document: {docx_path}")
    if pdf_path:
        print(f"PDF Document: {pdf_path}")
    print("="*60)
