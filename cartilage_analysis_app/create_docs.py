
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import re

# Analysis Methodology Text
TITLE = "Methodology: Automated Fiber Orientation Analysis of Articular Cartilage via Polarized Light Microscopy (PLM)"
SECTIONS = [
    {
        "title": "1. Image Acquisition and Preprocessing",
        "content": """The input data consists of high-resolution Polarized Light Microscopy (PLM) images of articular cartilage sections. The raw images, often in TIFF format, are first converted to standard RGB/PNG format for processing. A Region of Interest (ROI) is manually or semi-automatically cropped to isolate the cartilage depth profile, removing subchondral bone and background artifacts.

To ensure robust color analysis independent of lighting variations, the images are converted from the BGR (Blue-Green-Red) color space to the HSV (Hue-Saturation-Value) color space. The Hue (H) channel represents color independent of brightness, the Saturation (S) channel represents color purity, and the Value (V) channel represents light intensity."""
    },
    {
        "title": "2. Automated Zone Detection via K-Means Clustering and Spatial Voting",
        "content": """To overcome the limitations of manual zone segmentation, we implemented an unsupervised machine learning approach using K-Means clustering combined with spatial logic. This method robustly identifies the Superficial Zone (SZ), Middle Zone (MZ), and Deep Zone (DZ) based on their intrinsic color and intensity signatures.

2.1. Feature Extraction and Normalization
For every pixel in the ROI, the Hue (0-180) and Intensity (0-255) values are extracted and normalized to a 0-1 range. Pixels with low intensity (V < 10) are filtered out as background noise.

2.2. K-Means Clustering
A K-Means clustering algorithm (k=3) is applied to the normalized (Hue, Intensity) feature space. This classifies pixels into three distinct clusters based on their spectral properties:
   - Dark Extinction Cluster (MZ): Identified as the cluster with the lowest centroid intensity (V), corresponding to the "extinction" phenomenon observed in the Middle Zone.
   - Superficial Zone Cluster (SZ): Identified as the cluster with the lower mean Hue (typically Orange/Red spectrum).
   - Deep Zone Cluster (DZ): Identified as the cluster with the higher mean Hue (typically Green spectrum).

2.3. Spatial Row Voting
To enforce spatial continuity, a row-wise voting mechanism is employed. For each pixel row y in the image, the median properties (Hue, Intensity) are calculated. The row is then assigned to the zone (SZ, MZ, or DZ) whose cluster centroid is closest in Euclidean distance to the row's median properties. This generates a depth-dependent zone profile.

2.4. Boundary Definition
The transitions between zones are detected by scanning the simplified zone profile. The SZ/MZ boundary is defined as the depth where the classification consistently shifts from the SZ cluster to the MZ cluster, and the MZ/DZ boundary is defined similarly."""
    },
    {
        "title": "3. Quantitative Fiber Orientation Analysis",
        "content": """The fiber orientation angle (Theta) is derived quantitatively from the observed interference colors.

3.1. Color-to-Angle Mapping Calibration
The system dynamically calibrates the mapping function F(H) -> Angle for each image (or batch of images) based on the detected zones:
   - Zero-Degree Anchor (SZ Hue): The median hue of the detected Superficial Zone is defined as corresponding to fibers parallel to the surface (0 degrees).
   - Ninety-Degree Anchor (DZ Hue): The median hue of the detected Deep Zone is defined as corresponding to fibers perpendicular to the surface (90 degrees).

3.2. Pixel-wise Angle Calculation
For every valid pixel (x, y), the fiber orientation angle Theta(x, y) is calculated using a linear interpolation mapping function:
    Theta = ((Hue - SZ_Hue) / (DZ_Hue - SZ_Hue)) * 90

Values are clamped to the [0, 90] degree range. This method accounts for sample-specific color shifts (e.g., "Red-shifted" vs. "Green-shifted" samples) by using internal standards rather than hardcoded threshold values."""
    },
    {
        "title": "4. Depth Profiling and Statistical Aggregation",
        "content": """The morphological variation of fiber orientation with depth is quantified by generating a normalized depth profile.

1. Normalization: The cartilage thickness is normalized from 0.0 (Articular Surface) to 1.0 (Tidemark/Bone Interface).
2. Binning: The depth is discretized into 100 equally spaced bins (1% thickness increments).
3. Aggregation: For each depth bin, the orientation angles of all constituent pixels are aggregated. The circular mean orientation and standard deviation are calculated to account for the angular periodicity.
4. Gap Filling: To handle artifacts or voids (fissures, lacunae), a linear interpolation algorithm fills discontinuous data points in the depth profile, ensuring a continuous function for derivation of biomechanical gradients."""
    },
     {
        "title": "5. Batch Processing and Data Normalization",
        "content": """For multi-sample comparison, a Reference-Based Batch Analysis protocol is used. A representative "Reference Image" is first analyzed to establish the canonical calibration values. These calibration parameters are then globally applied to all subsequent samples in the batch. This ensures that quantitative comparisons (e.g., mean angular deviation, zonal thickness changes) reflect true structural differences rather than inter-sample staining or lighting inconsistencies."""
    }
]

OUT_DIR = "Methodology_Report"
if not os.path.exists(OUT_DIR):
    os.makedirs(OUT_DIR)

# --- Generate DOCX ---
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# Title
heading = doc.add_paragraph(TITLE)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading.runs[0].bold = True
heading.runs[0].font.size = Pt(16)

doc.add_paragraph() # Spacer

for sec in SECTIONS:
    h = doc.add_heading(sec['title'], level=1)
    # Styles for heading
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None # Default color
        run.font.size = Pt(14)
    
    # Content
    doc.add_paragraph(sec['content'])
    doc.add_paragraph()

docx_path = os.path.join(OUT_DIR, "Methodology.docx")
doc.save(docx_path)
print(f"Generated DOCX: {docx_path}")

# --- Generate PDF ---
pdf = FPDF()
pdf.add_page()
pdf.set_font("Times", size=12)

# Title
pdf.set_font("Times", 'B', 16)
pdf.multi_cell(0, 10, TITLE, align='C')
pdf.ln(10)

# Content
pdf.set_font("Times", size=12)
for sec in SECTIONS:
    # Heading
    pdf.set_font("Times", 'B', 14)
    pdf.cell(0, 10, sec['title'], ln=True)
    
    # Body
    pdf.set_font("Times", '', 12)
    # sanitize text for latin-1
    content = sec['content'].encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 6, content)
    pdf.ln(5)

pdf_path = os.path.join(OUT_DIR, "Methodology.pdf")
pdf.output(pdf_path)
print(f"Generated PDF: {pdf_path}")
